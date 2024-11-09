# -- coding: utf-8 --

import os
from PIL import Image, ImageTk
import tkinter as tk
import cv2
from PIL import Image, ImageTk
import google.generativeai as genai  #GEMINI-1 
from gtts import gTTS  #Texto a audio
from rich.console import Console #HTMLtexto consola
from rich.syntax import Syntax
import time


from io import BytesIO #almacenar img, audio 
import keyboard
import pygame
import pygame.camera
import re
import datetime


from docx import Document
import html2text 
import uuid

import subprocess
import glob
import datetime
import shutil


# Definición de la variable global
contador = 0           
class CameraApp: 

        
    def to_plain_string(self,text):
        # Eliminar caracteres especiales excepto letras, números y tildes
        cleaned_text = re.sub(r'[^\w\sáéíóúÁÉÍÓÚ.,;:?!¡¿%]', '', text)
        return cleaned_text 

    def contiene_palabra(self, cadena, palabra):
        # Verificar si la cadena no es None antes de intentar la comparación
        if cadena is not None:
            cadena = cadena.lower()  # Convertir la cadena a minúsculas
            palabra = palabra.lower()  # Convertir la palabra a minúsculas
            return palabra in cadena
        else:
            return False   

    def cerrar_ventana(self, event):
        self.master.destroy()  

    def hablar(self,texto):
        # Generar el archivo de audio usando gTTS y BytesIO
        print(texto)
        with BytesIO() as f:
            tts = gTTS(text=texto, lang='es', slow= False)
            tts.write_to_fp(f)
            f.seek(0)
            
            # Reproducir el audio usando pygame mixer
            pygame.mixer.init()
            pygame.mixer.music.load(f)
            pygame.mixer.music.play()
            
            # Esperar hasta que termine de reproducirse el audio
            while pygame.mixer.music.get_busy():
                if keyboard.is_pressed('Enter'):
                        pygame.mixer.music.stop()
                        time.sleep(0.1)
                        break
                continue
 

    def desconexion_audio(self):
        # Inicializar pygame
        pygame.init()

        try:
            # Iniciar la reproducción del audio

            pygame.mixer.music.load("desconexion.mpeg")
            pygame.mixer.music.play()

            # Esperar hasta que el audio termine de reproducirse
            while pygame.mixer.music.get_busy():
                pygame.time.Clock().tick(10)
        except pygame.error as e:
            print("Error al reproducir el audio:", e)
            self.error()
        finally:
            # Detener pygame
            pygame.quit()
 

    def guardado_exitoso(self):
        # Inicializar pygame
        pygame.init()

        try:
            # Iniciar la reproducción del audio
            pygame.mixer.music.load("guardado_exitoso.mpeg")
            pygame.mixer.music.play()

            # Esperar hasta que el audio termine de reproducirse
            while pygame.mixer.music.get_busy():
                pygame.time.Clock().tick(10)
        except pygame.error as e:
            print("Error al reproducir el audio:", e)
            self.error()
        finally:
            # Detener pygame
            pygame.quit() 
 

    def error(self):
        # Inicializar pygame
        pygame.init()

        try:
            # Iniciar la reproducción del audio
            pygame.mixer.music.load("error.mpeg")
            pygame.mixer.music.play()

            # Esperar hasta que el audio termine de reproducirse
            while pygame.mixer.music.get_busy():
                pygame.time.Clock().tick(10)
        except pygame.error as e:
            print("Error al reproducir el audio:", e)
            self.error()
        finally:
            # Detener pygame
            pygame.quit() 


    def html_to_docx(self,text):
                # Carpeta donde se guardarán los documentos
                folder_path = "respuestas"
                if not os.path.exists(folder_path):
                    os.makedirs(folder_path)  
                output_path = os.path.join(folder_path, "unl-text"+str(uuid.uuid4())+".docx")  
		 
                # Crear un nuevo documento de Word
                doc = Document()
                texto_sin_saltos = text.replace("\n", " ")
                text_content = html2text.html2text(texto_sin_saltos) 
                doc.add_paragraph(text_content)
                doc.save(output_path)
                print("Guardado en "+folder_path+" como: unl-text"+str(uuid.uuid4())+".docx")

                try:
                    usb_path = os.path.join("/media/unl/KINGSTON", "unl-text"+str(uuid.uuid4())+".docx") 
                    doc = Document()
                    texto_sin_saltos = text.replace("\n", " ")
                    text_content = html2text.html2text(texto_sin_saltos) 
                    doc.add_paragraph(text_content)
                    doc.save(usb_path)
                except Exception as e:
                    print("No se pudo guardar el archivo en un USB nombrado como Kingston")
                    #self.error()

#https://gtts.readthedocs.io/en/latest/cli.html
    def reproducir_audio(self,texto):
        try:
            with BytesIO() as f:
                tts = gTTS(text=texto_sin_saltos, lang='es')
                tts.write_to_fp(f)

                f.seek(0) 
                pygame.mixer.init()
                pygame.mixer.music.load(f)
                pygame.mixer.music.play()
                
                while pygame.mixer.music.get_busy():
                    if keyboard.is_pressed('0'):
                        pygame.mixer.music.stop()
                        time.sleep(0.1)
                        self.desconexion_audio()
                        break
                    
                    if keyboard.is_pressed('+'):
                        pos=pygame.mixer.music.get_pos()
                        pygame.mixer.music.set_pos(pos + 1) #adelantar un segundo 
                        
                    if keyboard.is_pressed('-'):
                        pos=pygame.mixer.music.get_pos()
                        pygame.mixer.music.set_pos(pos - 1) #adelantar un segundo 
                        
                    # Detectar si se ha pulsado la tecla "Enter"
                    pausado = False
                    if pausado == False:

                        #print("Esperando el . para pause")
                        if keyboard.is_pressed('Enter'):
                            pausado = True
                            pygame.mixer.music.pause()
                            #time.sleep(0.1)
                            print("pause")
                        elif keyboard.is_pressed('0'):
                            pygame.mixer.music.stop()
                            time.sleep(0.1) 
                            break 
                        while pausado == True:
                                print("Esperando el . para play")
                                if keyboard.is_pressed('Enter'):
                                    pausado==False 
                                    pygame.mixer.music.unpause()
                                    #time.sleep(0.1)
                                    print("Play") 
                                    break

                                elif keyboard.is_pressed('0'):
                                    pygame.mixer.music.stop()
                                    break
                    continue    
                self.preguntar_guardar(texto,tts) 
        except Exception as e:
            print("reproducir_audio(); GTTS ERROR - - ERROR - - ERROR   GTTS")
            self.error()
            print('{type(e).__name__}: '+ str(e))     

 


    def reproducir_audio_no_save(self,texto):
        try:
            with BytesIO() as f:
                texto_sin_saltos = texto.replace("\n", " ") 
                tts = gTTS(text=texto_sin_saltos, lang='es')
                tts.write_to_fp(f)

                f.seek(0) 
                pygame.mixer.init()
                pygame.mixer.music.load(f)
                pygame.mixer.music.play()
                
                while pygame.mixer.music.get_busy():
                    if keyboard.is_pressed('0'):
                        pygame.mixer.music.stop()
                        time.sleep(0.1)
                        self.desconexion_audio()
                        break
                    
                    if keyboard.is_pressed('+'):

                        pos=pygame.mixer.music.get_pos()
                        pygame.mixer.music.set_pos(pos + 1) #adelantar un segundo 
                        
                    if keyboard.is_pressed('-'):
                        pos=pygame.mixer.music.get_pos()
                        pygame.mixer.music.set_pos(pos - 1) #adelantar un segundo 
                        
                    # Detectar si se ha pulsado la tecla "Enter"
                    pausado = False
                    if pausado == False:

                        #print("Esperando el . para pause")
                        if keyboard.is_pressed('Enter'):
                            pausado = True
                            pygame.mixer.music.pause()
                            #time.sleep(0.1)
                            print("pause")
                        elif keyboard.is_pressed('0'):
                            pygame.mixer.music.stop()
                            time.sleep(0.1) 
                            break 
                        while pausado == True:

                                print("Esperando el . para play")
                                if keyboard.is_pressed('Enter'):
                                    pausado==False 
                                    pygame.mixer.music.unpause()
                                    #time.sleep(0.1)
                                    print("Play") 
                                    break

                                elif keyboard.is_pressed('0'):
                                    pygame.mixer.music.stop()
                                    break
                    continue     
        except Exception as e:
            print("reproducir_audio(); GTTS ERROR - - ERROR - - ERROR   GTTS")
            self.error()
            print('{type(e).__name__}: '+ str(e))     


#https://gtts.readthedocs.io/en/latest/cli.html
    def reproducir_audio_nosound(self,texto):
        try:
            with BytesIO() as f:
                texto_sin_saltos = texto.replace("\n", " ") 
                tts = gTTS(text=texto_sin_saltos, lang='es')
                tts.write_to_fp(f)

                f.seek(0) 
                pygame.mixer.init()
                pygame.mixer.music.load(f)
                pygame.mixer.music.play()
                
                while pygame.mixer.music.get_busy():
                    if keyboard.is_pressed('0'):
                        pygame.mixer.music.stop()
                        time.sleep(0.1) 
                        break
                    continue    

        except Exception as e:
            print("reproducir_audio(); GTTS ERROR - - ERROR - - ERROR   GTTS")
            self.error()
            print('{type(e).__name__}: '+ str(e))     

 

    def preguntar_guardar(self,texto,tts):
        try:
            self.reproducir_audio_nosound("Guardar texto tecla 7, audio 8, o ambos 9")
            tecla = keyboard.read_key()  
            if tecla=="7":
                self.html_to_docx(texto)    
                self.guardado_exitoso()
            if tecla=="8":
                tts = tts
                # Obtener la fecha y hora actual
                now = datetime.datetime.now()
                name = "respuestas/unl-audio-"+str(uuid.uuid4()) + ".mp3"

                # Carpeta donde se guardarán los audios 
                folder_path = "respuestas"
                if not os.path.exists(folder_path):
                    os.makedirs(folder_path)   
                tts.save(name)  
                
                try:
                    # USB donde se guardarán los audios   
                    nombre_audio_usb="/media/unl/KINGSTON/unl-audio-"+str(uuid.uuid4())+".mp3"
                    tts.save(nombre_audio_usb)
                    print("Guardado en KINGSTON")
                except Exception as e:              
                    try:
                        # USB donde se guardarán los audios   
                        nombre_audio_usb="/media/unl/kingston/unl-audio-"+str(uuid.uuid4())+".mp3"
                        tts.save(nombre_audio_usb)
                        print("Guardado en kingston")
                    except Exception as e:
                        print("No se guardó en kingston o KINGSTON")
                self.guardado_exitoso()
            if tecla=="9":
                self.html_to_docx(texto) 

                tts = tts
                # Obtener la fecha y hora actual
                now = datetime.datetime.now()

                # Carpeta donde se guardarán los audios 
                folder_path = "respuestas"
                if not os.path.exists(folder_path):
                    os.makedirs(folder_path)  
                nombre_audio = "respuestas/unl-audio-"+str(uuid.uuid4()) + ".mp3"
                tts.save(nombre_audio)
                try:
                    # USB donde se guardarán los audios   
                    nombre_audio_usb="/media/unl/KINGSTON/unl-audio-"+str(uuid.uuid4())+".mp3"
                    tts.save(nombre_audio_usb)
                    print("Guardado en KINGSTON")
                except Exception as e:              
                    try:
                        # USB donde se guardarán los audios   
                        nombre_audio_usb="/media/unl/kingston/unl-audio-"+str(uuid.uuid4())+".mp3"
                        tts.save(nombre_audio_usb)
                        print("Guardado en kingston")
                    except Exception as e:

                        print("No se guardó en kingston o KINGSTON")

                self.guardado_exitoso()

            elif not tecla=="0":
                exit 
        
        except Exception as e:
            print("Error en la guardada del archivo")
            self.error()
            print('{type(e).__name__}: '+ str(e))     

 




    def gemini_unl(self, mensaje):
                foto, path =self.tomarFoto() 
                #Eliminar Foto del repositorio 
                #os.remove(path)  
                print(" mensaje ",mensaje)  
                print(" foto ",foto)    
                model = genai.GenerativeModel('gemini-1.5-flash')  
                # Seguridad https://ai.google.dev/docs/safety_setting_gemini?hl=es-419

                safety_settings = {
                    "HARM_CATEGORY_HARASSMENT": "BLOCK_NONE",
                    "HARM_CATEGORY_HATE_SPEECH": "BLOCK_NONE",
                    "HARM_CATEGORY_SEXUALLY_EXPLICIT": "BLOCK_NONE",
                    "HARM_CATEGORY_DANGEROUS_CONTENT": "BLOCK_NONE"} 
                try:
                    response = model.generate_content([mensaje, foto], safety_settings=safety_settings)
                    #time.sleep(1)  
                    #time.sleep(1)  
                    response.resolve() 
                    print(response)  
                    if response.candidates[0].content.parts: #Pa ver si hay más de una respuesta:
                        try: 
                            respuesta= response.candidates[0].content.parts[0].text   
                            respuesta= self.to_plain_string(respuesta) 
                            #Devolviendo respuesta
                            print(respuesta) 
                            return(respuesta)  
                        except Exception as e:
                            print("Entró al Primer CATCH, ERROR - - ERROR - - ERROR")
                            print('f{type(e).__name__}: ',{e})     
                            try:  # Si solo hay una respuesta: 
                                print("!!!!!!!!!!!!!!!")
                                a=response.text    
                                respuestab=self.to_plain_string(a)
                                #Devolviendo respuesta
                                return(respuestab)  

                            except Exception as e:
                                print("Entró al segundo CATCH, ERROR - - ERROR - - ERROR")
                                self.hablar("No se pudo generar una respuesta")
                                print('{type(e).__name__}: '+{e})     

                    else: # Si solo hay una respuesta:
                        try:  
                            print("Parece que solo hay una respuesta. :)")
                            print("!!!!!!!!!!!!!!!")
                            a=response.text   
                            print("response.text")
                            respuesta=self.to_plain_string(respuesta)  
                            #Devolviendo respuesta
                            return(respuesta)  

                        except Exception as e:
                            print("Entró al segundo CATCH, del Else")
                            print('{type(e).__name__}: '+{e})     
                            try:  # Si solo hay una respuesta:
                                print("Parece que solo hay una respuesta. :)")  
                                respuesta= response.candidates[0].content.parts[0].text 
                                #Devolviendo respuesta
                                return(respuesta)  

                            except Exception as e:
                                print("ERROR - - ERROR")
                                self.error()
                                print('f{type(e).__name__}: ',+str(e))     
                    
                except Exception as e:
                    print("No se pudo generar una respuesta")
                    self.error() 
                    print('f{type(e).__name__}: ',{e})    
                

#####
    def gemini(self, prompt):
            self.start_function()  
            prompt=prompt
            if self.contiene_palabra(prompt, "Identifica"):   
                print("Ejecutar sin guardar")
                try:  
                    res = self.gemini_unl(prompt) 
                    #REPRODUCIR 
                    self.reproducir_audio_no_save(res)  
                    self.desconexion_audio()
                    self.executing = False
                    self.finish_function()  

                except Exception as e:
                    print("Hubo un error en la ultima funcion!")
                    self.error()
                    self.executing = False
                    self.finish_function()  
                    print('{type(e).__name__}: '+str(e))    
                
            else: 
                print("Ejecutar y guardar") 
                try:  
                    res = self.gemini_unl(prompt)
                    resnospaces=self.eliminar_saltos_de_linea(res)
                    #REPRODUCIR 
                    self.reproducir_audio(resnospaces)  
                    self.desconexion_audio()
                    self.executing = False
                    self.finish_function()  

                except Exception as e:
                    print("Hubo un error en la ultima funcion!")
                    self.error()
                    self.executing = False
                    self.finish_function()  
                    print('{type(e).__name__}: '+str(e))  
########



    def __init__(self, window, window_title):  

            GOOGLE_API_KEY='AIzaSyCI5hk_LHr8V49BjaI4pm6BxpZBwJplxII'
            genai.configure(api_key=GOOGLE_API_KEY) 
            self.window = window
            self.window.title(window_title)

            #Superponer la ventana
            #self.window.attributes('-topmost',True)    
 

                # IMAGEN BACKGROUND
            image_path = ImageTk.PhotoImage(file="img.jpeg")
 
            self.photo = image_path

            # Crear un widget Label para mostrar la imagen
            self.image_label = tk.Label(window, image=self.photo)
            self.image_label.pack()

            # Mover la imagen a la izquierda (coordenadas x=0)
            self.image_label.place(x=0, y=0) 

            #Actualizar la camara
            #self.update_video() 
  

## COMANDS  >>> 
            self.executing = False
            self.window.bind('<Escape>', lambda e: self.exit()) 
            self.window.bind('<KeyPress>', self.on_key_press) 
            self.window.mainloop()

 
   
    def on_key_press(self, event):  
        if not self.executing:
            if event.keysym:  
                self.executing = True
                if event.keysym == '0' or event.keysym == 'KP_0':
                    print("Tecla '0' presionada")
                    self.executing = False
                elif event.keysym == '1' or event.keysym == 'KP_1':
                    print("Tecla '1' presionada")
                    self.gemini("Identifica que hay en la imagen, hay texto?, tablas? graficos? en español: solo responde algo como: el contenido del documento es texto, el contenido del documento es gráficos, o el contenido del documento es tablas, o el contenido del documento es texto y tablas o texto y gráficos, o si hay texto tablas y gráficos, o no hay, nada, dependiendo del contenido del documento, no des una explicación del documento.")
                    self.executing = False
                elif event.keysym == '2' or event.keysym == 'KP_2':
                    print("Tecla '2' presionada")
                    self.gemini("extrae el texto que encuentres")
                    self.executing = False
                elif event.keysym == '3' or event.keysym == 'KP_3':
                    print("Tecla '3' presionada")
                    self.gemini("Narra en español todo el contenido de la tabla de una forma ordenada, coherente y estructurada de forma que se pueda entender TODO el contenido, primero narra las columnas y luego narra los valores de cada fila relacionando con cada columna")
                    self.executing = False
                elif event.keysym == '4' or event.keysym == 'KP_4':
                    print("Tecla '4' presionada")
                    self.gemini("Describe todo el contenido del gráfico en español, y leelo si fuera necesario, de una forma ordenada, coherente y estructurada de forma que se pueda entender TODO el contenido")
                    self.executing = False
                elif event.keysym == '5' or event.keysym == 'KP_5':
                    print("Tecla '5' presionada")
                    self.gemini("Extrae el texto de los párrafos y Describe todo el contenido, si hay tablas o gráfico, describe cada fila o sección explicando de forma ordenada, coherente y estructurada de forma que se pueda entender TODO el contenido")
                    self.executing = False
                elif event.keysym == '6' or event.keysym == 'KP_6':
                    print("Tecla '6' presionada")
                    self.gemini("Describe en español detalladamente la imagen ")
                    self.executing = False
                elif event.keysym == '7' or event.keysym == 'KP_7':
                    print("Tecla '7' presionada")
                    self.executing = False
                elif event.keysym == '8' or event.keysym == 'KP_8':
                    print("Tecla '8' presionada")
                    self.executing = False
                elif event.keysym == '9' or event.keysym == 'KP_9':
                    print("Tecla '9' presionada")
                    self.executing = False
                elif event.keysym == '*' or event.keysym == 'KP_Multiply':
                    print("Tecla '*' presionada")
                    self.incrementar_contador()
                    self.info() 
                    self.executing = False
                elif event.keysym == '/' or event.keysym == 'KP_Divide':
                    print("Tecla '/' presionada")
                    self.decrementar_contador()
                    self.info() 
                    self.executing = False

## COMANDS <<< 



    def tomarFoto(self):
        self.subprocess()
        # Busca todos los archivos JPG en el directorio
        archivos_jpg = glob.glob(os.path.join(".", "*.jpg"))

        # Ordena los archivos por fecha de creación
        archivos_jpg.sort(key=os.path.getmtime)
        
        # Obtiene el último archivo creado
        ultimo_archivo_jpg = archivos_jpg[-1]
        
        # Abre el último archivo JPG
        imagen = Image.open(ultimo_archivo_jpg)  
        
        #os.remove(ultimo_archivo_jpg)
         
        return (imagen,ultimo_archivo_jpg) 
    def incrementar_contador(self):
        # Indicar que se usará la variable global 'contador'
        global contador
        contador += 1
        print(f'El contador es ahora: {contador}')

    def decrementar_contador(self):
        # Indicar que se usará la variable global 'contador'
        global contador
        contador = contador-1
        print(f'El contador es ahora: {contador}')

    def resetear_contador(self):
        global contador
        contador = 0
        print('El contador ha sido reiniciado')
 
    def info(self):
        global contador
        self.start_function()
        if contador == 1:
            self.reproducir_audio_no_save("Bienvenido a Eye Reader")
            self.finish_function()
        elif contador == 2:
            self.reproducir_audio_no_save("Para iniciar, coloca un documento en la bandeja")
            self.finish_function()
        elif contador == 3:
            self.reproducir_audio_no_save("Con la tecla 1 identifica el tipo de contenido")
            self.finish_function()
        elif contador == 4:
            self.reproducir_audio_no_save("La tecla 2 ejecuta la lectura de texto")
            self.finish_function()
        elif contador == 5:
            self.reproducir_audio_no_save("La tecla 3 para describir una tabla")
            self.finish_function()
        elif contador == 6:
            self.reproducir_audio_no_save("La tecla 4 para describir un gráfico")
            self.finish_function()
        elif contador == 7:
            self.reproducir_audio_no_save("La tecla 5 para leer texto con tablas y/o gráficos")
            self.finish_function()
        elif contador == 8:
            self.reproducir_audio_no_save("Puedes guardar el texto o audio generado; Despues de cada ejecución presiona 7 para guardar el texto, 8 para guardar el audio o 9 para guaradr ambos")
            self.finish_function()
        else:
            pass 

    def start_function(self):
        # Bloquear el teclado
        self.is_blocked = True
        print("Teclado bloqueado")

    def finish_function(self):
        # Termina la función y desbloquea el teclado
        self.is_blocked = False
        print("Teclado desbloqueado")


    def subprocess(self):
        subprocess.run(["nvgstcapture-1.0 --automate --capture-auto"], shell=True, capture_output=False, text=False) 

    def eliminar_saltos_de_linea(self, texto):
        texto_sin_saltos = texto.replace('\n', '')
        texto_limpio = re.sub(r'\s+', ' ', texto_sin_saltos).strip()
        return texto_sin_saltos 
 
    def exit(self): 
        self.window.destroy()

App = CameraApp(tk.Tk(), "Lector de documentos - UNL")
